"""
strategy_loader.py — Loads and structures the chesco funding strategy document.

Reads the Word document in data/ and returns a formatted string with headings
preserved, suitable for use as LLM context.
"""

import os
import re
from pathlib import Path

_STRATEGY_DIR = Path(__file__).parent / "data"
_STRATEGY_FILENAME = "Chesco Forschungsstrategie 230326v1 (2).docx"
_STRATEGY_PATH = _STRATEGY_DIR / _STRATEGY_FILENAME

_HEADING_STYLES = {"Heading 1", "Heading 2", "Heading 3", "Heading 4"}

_cached_strategy: str | None = None


def load_strategy_text() -> str:
    """
    Extracts full text from the chesco strategy Word document.

    Headings are prefixed with markdown-style # markers so LLMs can parse
    document structure. Returns a single formatted string.

    Raises FileNotFoundError if the strategy document is missing.
    Raises ImportError if python-docx is not installed.
    """
    global _cached_strategy
    if _cached_strategy is not None:
        return _cached_strategy

    if not _STRATEGY_PATH.exists():
        raise FileNotFoundError(
            f"Strategy document not found at: {_STRATEGY_PATH}\n"
            "Place the chesco strategy .docx file in the data/ directory."
        )

    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError(
            "python-docx is required to read the strategy document. "
            "Install it with: pip install python-docx"
        ) from exc

    doc = Document(str(_STRATEGY_PATH))
    lines: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style = para.style.name if para.style else "Normal"

        if style == "Heading 1":
            lines.append(f"\n# {text}")
        elif style == "Heading 2":
            lines.append(f"\n## {text}")
        elif style in ("Heading 3", "Heading 4"):
            lines.append(f"\n### {text}")
        else:
            lines.append(text)

    _cached_strategy = "\n".join(lines).strip()
    return _cached_strategy


_TARGET_SECTOR_PATTERN = re.compile(r"^4\.\d+\s+(.+)$")

_TECHNOLOGY_TERMS = [
    "Digital Twin", "Digital Twins",
    "Digitaler Zwilling", "Digitale Zwillinge", "digitalen Zwilling",
    "Digital Thread", "Digital-Thread",
    "Asset Administration Shell",
    "AAS", "OPC UA", "MQTT",
    "Manufacturing-X", "Manufacturing X",
    "Gaia-X", "Gaia X",
    "IDSA",
    "Co-Bot", "Cobot",
    "VR", "AR", "IEC 61508", "IEC 62443",
    "EU Cyber Resilience Act", "EU CRA", "Cyber Resilience Act",
    "AS9100", "AS 9100", "EASA", "AQAP",
    "Compliance by Design", "Compliance-by-Design",
]

_FOCUS_AREA_TERMS = [
    "Datensouveränität", "data sovereignty",
    "Digital Thread", "Digital-Thread",
    "Digital Twin", "Digital Twins",
    "Digitaler Zwilling", "Digitale Zwillinge", "digitalen Zwilling",
    "Asset Administration Shell",
    "Manufacturing-X", "Manufacturing X",
    "Compliance by Design", "Compliance-by-Design",
    "föderierte Datenarchitektur", "föderierten Datenarchitektur",
    "Cybersecurity-by-Design", "Cybersecurity by Design",
]

# Call-side vocabulary that maps to each doc sector; not literal doc text —
# funding calls phrase things differently than the strategy doc's section titles.
_SECTOR_SYNONYMS = {
    # Bare "raumfahrt" deliberately excluded -- it collides with the German
    # research ministry's own name ("Bundesministerium für Forschung,
    # Technologie und Raumfahrt (BMFTR)"), which appears in every BMFTR call
    # regardless of topic. Use more specific space-related terms instead.
    "Luft- und Raumfahrt": [
        "luftfahrt", "aerospace", "aviation", "drohne", "drone", "easa", "lufo",
        "raumfahrtprogramm", "weltraum", "satelliten",
    ],
    "Fertigung / Industrie 4.0": ["fertigung", "industrie 4.0", "manufacturing", "produktion", "montage"],
    "Zivile Sicherheit": ["zivile sicherheit", "civil security", "sifo", "feuerwehr", "kritische infrastruktur", "polizei", "rettungsdienst"],
}


def load_strategy_structure() -> dict:
    """
    Extracts structured strategy elements from the strategy document text:
    target_sectors (chesco's priority sectors, each with sector-local
    technology keywords found in that sector's own text), technologies
    (named tech/standards mentioned anywhere in the doc), focus_areas
    (core positioning themes), and exclusion_criteria (currently empty --
    the doc states priorities, not explicit no-go domains).
    """
    text = load_strategy_text()

    return {
        "target_sectors": _extract_target_sectors(text),
        "technologies": [t for t in _TECHNOLOGY_TERMS if t.lower() in text.lower()],
        "focus_areas": [t for t in _FOCUS_AREA_TERMS if t.lower() in text.lower()],
        "exclusion_criteria": [],
    }


def _extract_target_sectors(text: str) -> list[dict]:
    sectors: list[dict] = []
    current: dict | None = None

    for raw_line in text.split("\n"):
        line = raw_line.strip()
        is_heading = line.startswith("#")
        plain = line.lstrip("#").strip()

        sector_match = _TARGET_SECTOR_PATTERN.match(plain)
        if sector_match:
            if current:
                sectors.append(current)
            current = {"name": sector_match.group(1), "_body": ""}
            continue

        if is_heading:
            if current:
                sectors.append(current)
            current = None
            continue

        if current is not None and plain:
            current["_body"] += " " + plain

    if current:
        sectors.append(current)

    for sector in sectors:
        body = sector.pop("_body")
        sector["keywords"] = [t for t in _TECHNOLOGY_TERMS if t.lower() in body.lower()]
        sector["synonyms"] = _SECTOR_SYNONYMS.get(sector["name"], [])

    return sectors


# ---------------------------------------------------------------------------
# Table extraction — doc.paragraphs (used above) does not include table cell
# text at all, so the 4 tables in the strategy doc (capability/tools map,
# challenge-to-funding-body map, capability/maturity roadmap, funder/program/
# topic directory) were previously invisible to this loader.
# ---------------------------------------------------------------------------

_TABLE_LABELS = [
    "capability_tools",       # Table 0: Digitale Bausteine -> Fähigkeit -> Tools
    "challenge_funding_body", # Table 1: Ebene -> Herausforderung -> Fördermittelgeber
    "capability_maturity",    # Table 2: Fähigkeit -> Reifegrad -> Entwicklungsrichtung
    "funder_program_topic",   # Table 3: Fördergeber -> Programm -> Themen
]

# German/English synonym bridge for capability_maturity's "Fähigkeit" names --
# the table mixes English capability names (e.g. "Digital Twin") with German
# call-text vocabulary (e.g. "Digitale Zwillinge") that won't otherwise match
# via simple separator normalization. Only pairs actually observed to diverge
# between Fähigkeit names and existing _TECHNOLOGY_TERMS/_FOCUS_AREA_TERMS are
# listed here -- other Fähigkeit names (AAS-Implementierung, Co-Bot-Interface,
# Cybersecurity-by-Design, VR/AR-Schnittstellen) already match directly.
_CAPABILITY_SYNONYMS = {
    "Digital Twin": ["Digitaler Zwilling", "Digitale Zwillinge", "digitalen Zwilling", "Digital Twins"],
}

_cached_tables: dict | None = None


def _is_vertical_merge_continuation(cell) -> bool:
    """
    True if `cell` is a continuation of a vertically-merged cell above it
    (its own text is empty by construction; the real value lives in the
    cell that started the merge).
    """
    from docx.oxml.ns import qn

    tc_pr = cell._tc.tcPr
    if tc_pr is None:
        return False
    v_merge = tc_pr.find(qn("w:vMerge"))
    if v_merge is None:
        return False
    val = v_merge.get(qn("w:val"))
    return val is None or val == "continue"


def _parse_table(table) -> list[dict]:
    """
    Parses a docx table into a list of row-dicts keyed by the header row's
    cell text. python-docx's `row.cells` always has one entry per grid
    column (horizontally-merged cells simply repeat the same cell at each
    spanned position, which keeps header/value alignment correct without
    any special-casing). Vertically-merged continuation cells render as
    empty text in python-docx -- rather than silently guessing the parent
    value, those are flagged explicitly as "<merged-up>" so it's visible
    in the raw output which cells are real vs. merge artifacts.
    """
    rows = list(table.rows)
    if not rows:
        return []

    headers = [_clean_line(cell.text) for cell in rows[0].cells]

    parsed_rows = []
    for row in rows[1:]:
        values = []
        for cell in row.cells:
            if _is_vertical_merge_continuation(cell):
                values.append("<merged-up>")
            else:
                values.append(_clean_line(cell.text))
        parsed_rows.append(dict(zip(headers, values)))

    return parsed_rows


def _clean_line(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def load_strategy_tables() -> dict:
    """
    Extracts the 4 tables in the strategy document that load_strategy_text()
    cannot see (python-docx's doc.paragraphs excludes table cell content).
    Returns a dict keyed by _TABLE_LABELS, each value a list of row-dicts
    keyed by that table's header row.

    Raises FileNotFoundError if the strategy document is missing.
    Raises ImportError if python-docx is not installed.
    """
    global _cached_tables
    if _cached_tables is not None:
        return _cached_tables

    if not _STRATEGY_PATH.exists():
        raise FileNotFoundError(
            f"Strategy document not found at: {_STRATEGY_PATH}\n"
            "Place the chesco strategy .docx file in the data/ directory."
        )

    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError(
            "python-docx is required to read the strategy document. "
            "Install it with: pip install python-docx"
        ) from exc

    doc = Document(str(_STRATEGY_PATH))

    _cached_tables = {
        label: _parse_table(table)
        for label, table in zip(_TABLE_LABELS, doc.tables)
    }

    for row in _cached_tables.get("capability_maturity", []):
        row["synonyms"] = _CAPABILITY_SYNONYMS.get(row.get("Fähigkeit", ""), [])

    return _cached_tables
